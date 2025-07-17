import os
import codecs
import logging
from typing import List, Tuple, Dict, Any

# --- 1. 修改/新增导入 ---
from dotenv import load_dotenv
from pypdf import PdfReader
import docx

# 使用新版本的导入方式
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import ZhipuAIEmbeddings
from langchain.text_splitter import SpacyTextSplitter
from langchain.docstore.document import Document
# vvvvvvvvvvvvvvvv 核心修改点 vvvvvvvvvvvvvvvv
from chromadb.config import Settings # <-- 新增：导入Chroma的设置类
# ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

# --- 日志配置 ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class RAGIngestor:
    """
    一个用于处理文档、生成嵌入并将其存储到ChromaDB的类。
    核心功能是支持增量更新，避免重复处理已存在的文件。
    """

    def __init__(self, persist_dir: str, embeddings_model: str = "embedding-2"):
        """
        初始化Ingestor。
        :param persist_dir: ChromaDB持久化存储的路径。
        :param embeddings_model: 使用的嵌入模型名称。
        """
        self.persist_dir = persist_dir
        self.embeddings = ZhipuAIEmbeddings(model=embeddings_model)

        logging.info("正在初始化Spacy文本分割器 (使用 'xx_sent_ud_sm' 多语言模型)...")
        self.text_splitter = SpacyTextSplitter(
            pipeline="xx_sent_ud_sm",
            chunk_size=1000,
            chunk_overlap=200
        )
        logging.info("Spacy文本分割器初始化完成。")

        # vvvvvvvvvvvvvvvv 核心修改点 vvvvvvvvvvvvvvvv
        # --- 新增：定义客户端设置以禁用遥测 ---
        # 这是解决 SSLError 的关键，它会阻止Chroma尝试连接到 posthog.com
        self.client_settings = Settings(
            anonymized_telemetry=False,
            is_persistent=True,
        )
        # ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

        self.vectordb = self._load_or_create_vectordb()

    def _load_or_create_vectordb(self) -> Chroma:
        """加载现有的ChromaDB，如果不存在则创建一个新的。"""
        if os.path.exists(self.persist_dir):
            logging.info(f"从 {self.persist_dir} 加载已存在的向量数据库...")
            # --- 修改：在初始化时传入 client_settings ---
            return Chroma(
                persist_directory=self.persist_dir,
                embedding_function=self.embeddings,
                client_settings=self.client_settings # <-- 应用设置
            )
        else:
            logging.info("未找到现有数据库，将创建一个新的。")
            # --- 修改：在初始化时传入 client_settings ---
            return Chroma(
                embedding_function=self.embeddings,
                persist_directory=self.persist_dir,
                client_settings=self.client_settings # <-- 应用设置
            )

    def _get_ingested_files(self) -> set:
        """从向量数据库中获取所有已经处理过的文件名。"""
        if not self.vectordb or self.vectordb._collection.count() == 0:
            return set()

        try:
            existing_docs = self.vectordb.get(include=["metadatas"])
            ingested_sources = {meta['source'] for meta in existing_docs['metadatas']}
            logging.info(f"数据库中已存在 {len(ingested_sources)} 个来源文件。")
            return ingested_sources
        except Exception as e:
            logging.error(f"从数据库获取元数据失败: {e}")
            return set()

    def _read_supported_files(self, directory: str, ingested_files: set) -> List[Document]:
        """读取目录中所有支持的、且尚未被处理的新文件（.txt, .pdf, .docx）。"""
        documents = []
        supported_extensions = [".txt", ".pdf", ".docx"]
        logging.info(f"开始扫描目录 {directory} 中的新文件: {', '.join(supported_extensions)}")

        for filename in os.listdir(directory):
            if filename in ingested_files:
                continue

            file_path = os.path.join(directory, filename)
            file_extension = os.path.splitext(filename)[1].lower()

            if file_extension not in supported_extensions:
                continue

            content = ""
            try:
                if file_extension == ".txt":
                    with codecs.open(file_path, 'r', encoding='utf-8') as file:
                        content = file.read()
                elif file_extension == ".pdf":
                    reader = PdfReader(file_path)
                    text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
                    content = "\n".join(text_parts)
                elif file_extension == ".docx":
                    doc = docx.Document(file_path)
                    text_parts = [para.text for para in doc.paragraphs if para.text]
                    content = "\n".join(text_parts)

                if content.strip():
                    metadata = {"source": filename}
                    documents.append(Document(page_content=content, metadata=metadata))
                    logging.info(f"发现并读取新文件: {filename}")
                else:
                    logging.warning(f"文件 {filename} 为空或未能提取到文本内容，已跳过。")

            except Exception as e:
                logging.error(f"读取文件 {filename} 失败: {str(e)}")
        return documents

    def ingest_from_directory(self, directory: str):
        """
        从指定目录执行增量式的数据摄入，并使用分批处理来避免API限制。
        """
        ingested_files = self._get_ingested_files()
        new_documents = self._read_supported_files(directory, ingested_files)

        if not new_documents:
            logging.info("没有发现需要处理的新文件。")
            return

        logging.info(f"正在使用Spacy分割 {len(new_documents)} 个新文档...")
        split_docs = self.text_splitter.split_documents(new_documents)
        logging.info(f"文档被分割成 {len(split_docs)} 个块。")

        # 优化：为每个块添加更丰富的元数据
        for doc in split_docs:
            source = doc.metadata["source"]
            if not hasattr(self, f'chunk_count_{source}'):
                setattr(self, f'chunk_count_{source}', 0)

            chunk_num = getattr(self, f'chunk_count_{source}')
            doc.metadata['chunk_number'] = chunk_num
            setattr(self, f'chunk_count_{source}', chunk_num + 1)

        # 将新文档分批添加到向量数据库
        batch_size = 64
        total_docs = len(split_docs)
        logging.info(f"开始分批添加文档到ChromaDB，每批最多 {batch_size} 个。")

        for i in range(0, total_docs, batch_size):
            batch = split_docs[i:i + batch_size]
            logging.info(
                f"正在处理批次 {i // batch_size + 1}/{(total_docs + batch_size - 1) // batch_size}，包含 {len(batch)} 个文档块...")

            try:
                self.vectordb.add_documents(batch)
                logging.info(f"批次 {i // batch_size + 1} 添加成功。")
            except Exception as e:
                logging.error(f"处理批次 {i // batch_size + 1} 时发生错误: {e}")
                continue

        # 在所有批次处理完毕后，统一持久化
        logging.info("所有批次处理完毕，正在持久化数据库...")
        # 注意：新版本的Chroma会自动持久化，但显式调用.persist()在旧版本中是必要的，保留也无害。
        self.vectordb.persist()
        logging.info(f"数据摄入完成！向量数据库已更新并保存至: {os.path.abspath(self.persist_dir)}")


# --- 使用示例 ---
if __name__ == "__main__":
    load_dotenv()

    if "ZHIPUAI_API_KEY" not in os.environ:
        raise ValueError("请在您的 .env 文件或环境变量中设置 ZHIPUAI_API_KEY")

    script_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(script_dir, "LLM1_ChromaDB_Spacy")
    source_directory = r'C:\Users\lhy\Desktop\graph'

    ingestor = RAGIngestor(persist_dir=db_path)
    ingestor.ingest_from_directory(directory=source_directory)

    print("\n--- RAG数据摄入流程执行完毕 ---")

    if ingestor.vectordb:
        print(f"数据库中总文档块数: {ingestor.vectordb._collection.count()}")
        try:
            results = ingestor.vectordb.similarity_search("孙悟空有什么法宝？", k=2)
            print("\n示例查询结果:")
            for doc in results:
                print(f"来源: {doc.metadata.get('source')}, 块号: {doc.metadata.get('chunk_number')}")
                print(f"内容: {doc.page_content[:100]}...\n")
        except Exception as e:
            print(f"执行示例查询失败: {e}")